import os
import re
import base64
import urllib.request
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# 6个 Mermaid 图表的规范图注映射
MERMAID_TITLES = {
    1: "图 4-1 系统总体分层与架构设计图",
    2: "图 4-2 系统数据库 E-R 关系图",
    3: "图 4-3 用户登录鉴权业务时序图",
    4: "图 4-4 设备领用审批业务时序图",
    5: "图 4-5 设备检修闭环工作流时序图",
    6: "图 4-6 AI 运营报告生成时序图"
}

def set_cell_background(cell, hex_color):
    """设置单元格背景颜色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """设置单元格内边距"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_run_font(run, font_name='宋体', font_size_pt=12, bold=False, italic=False, color=None):
    """精确设置 Run 的中英文字体、字号及样式"""
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(font_size_pt)
    
    # 设置西文字体
    run.font.name = 'Times New Roman'
    # 设置中文字体 (eastAsia)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def format_paragraph(paragraph, line_spacing=1.5, space_before_pt=0, space_after_pt=6, first_line_indent_chars=0):
    """精确设置段落格式，包括行距、段前段后间距和首行缩进"""
    p_format = paragraph.paragraph_format
    p_format.line_spacing = line_spacing
    p_format.space_before = Pt(space_before_pt)
    p_format.space_after = Pt(space_after_pt)
    if first_line_indent_chars > 0:
        p_format.first_line_indent = Pt(12 * first_line_indent_chars)
    else:
        p_format.first_line_indent = Pt(0)

def set_table_borders(table):
    """为表格应用细灰色网格线"""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # 细灰色边框设置
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 0.5 磅
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3') # 浅灰色
        tblBorders.append(border)
    tblPr.append(tblBorders)

def autofit_table_widths(table, total_width_cm=15.5):
    """根据内容字符长度自适应调整表格列宽"""
    num_cols = len(table.columns)
    if num_cols == 0:
        return
    
    # 统计每列的最大字符数（中文字符算2个长度，英文字符/半角算1个长度）
    max_lens = [0] * num_cols
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            text = cell.text
            length = 0
            for char in text:
                if '\u4e00' <= char <= '\u9fff':
                    length += 2
                else:
                    length += 1
            if length > max_lens[c_idx]:
                max_lens[c_idx] = length
                
    # 设定一个最小参考长度
    for idx in range(num_cols):
        if max_lens[idx] < 4:
            max_lens[idx] = 4
            
    total_len = sum(max_lens)
    
    # 按照比例分配宽度，并进行安全限制（每列至少 1.2 cm）
    min_width = 1.2
    remaining_width = total_width_cm - (num_cols * min_width)
    
    col_widths = []
    if remaining_width < 0:
        col_width_avg = total_width_cm / num_cols
        col_widths = [col_width_avg] * num_cols
    else:
        for idx in range(num_cols):
            w = min_width + (max_lens[idx] / total_len) * remaining_width
            col_widths.append(w)
            
    # 应用到列和单元格
    for c_idx, col in enumerate(table.columns):
        col.width = Cm(col_widths[c_idx])
        for cell in col.cells:
            cell.width = Cm(col_widths[c_idx])

def write_table_to_docx(doc, table_rows):
    """把缓存的表格行数据渲染成 Word 表格"""
    if not table_rows:
        return
    num_rows = len(table_rows)
    num_cols = max(len(r) for r in table_rows)
    
    # 创建表格
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    
    # 逐行填入数据
    for r_idx, row_data in enumerate(table_rows):
        for c_idx, cell_value in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
            
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            
            # 智能对齐设置
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                val_stripped = cell_value.strip()
                # 长度 <= 6 个字符，或是一些勾选标识，或者是测试用例编号
                if len(val_stripped) <= 6 or val_stripped in ["√", "×"] or re.match(r'^TC-[A-Z]+-\d+$', val_stripped):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 表头特殊处理
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2") # 浅灰色背景表头
                run = p.add_run(cell_value)
                set_run_font(run, font_name='黑体', font_size_pt=10.5, bold=True)
            else:
                # 内容行
                # 斑马纹交替底色
                if r_idx % 2 == 1:
                    set_cell_background(cell, "FCFCFC")
                run = p.add_run(cell_value)
                set_run_font(run, font_name='宋体', font_size_pt=10, bold=False)
                
    # 避免行跨页撕裂
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
    # 设置表头在下一页重复显示
    header_tr = tbl.rows[0]._tr.get_or_add_trPr()
    header_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    # 自动调整列宽
    autofit_table_widths(tbl)
    
    # 表格后加空行
    p_spacer = doc.add_paragraph()
    format_paragraph(p_spacer, line_spacing=1.0, space_before_pt=0, space_after_pt=6)

def write_code_block_to_docx(doc, code_lines, code_lang):
    """把缓存的代码行数据渲染成带浅灰背景色的 Word 代码框"""
    if not code_lines:
        return
    # 创建表格
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Cm(14.5)
    set_cell_background(cell, "F5F5F5") # 浅灰色背景
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    
    code_content = "".join(code_lines).rstrip('\n')
    run = p.add_run(code_content)
    set_run_font(run, font_name='宋体', font_size_pt=9.5, bold=False, color=RGBColor(0x33, 0x33, 0x33))
    
    # 添加细灰色框线
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'E0E0E0')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    # 后面加一个空的微小段落作为间距
    p_spacer = doc.add_paragraph()
    format_paragraph(p_spacer, line_spacing=1.0, space_before_pt=0, space_after_pt=4)

def add_table_of_contents(paragraph):
    """在段落中添加 Word 自动目录域"""
    run = paragraph.add_run()
    # 域定义开始
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def enable_update_fields_on_open(doc):
    """开启设置：Word 打开时提示自动更新所有域（如目录）"""
    settings = doc.settings._element
    updateFields = OxmlElement('w:updateFields')
    updateFields.set(qn('w:val'), 'true')
    settings.append(updateFields)

def download_mermaid_image(mermaid_code, output_path):
    """调用 kroki.io 接口，将 Mermaid 代码在线渲染并下载为本地 PNG 图片（支持中文且更稳定）"""
    import zlib
    import base64
    import urllib.request
    import time
    
    # 对代码进行 UTF-8 编码并使用 zlib 压缩
    code_bytes = mermaid_code.encode('utf-8')
    compressed = zlib.compress(code_bytes, 9)
    # 使用 URL 安全的 Base64 编码
    encoded_string = base64.urlsafe_b64encode(compressed).decode('utf-8')
    url = f"https://kroki.io/mermaid/png/{encoded_string}"
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            req = urllib.request.Request(
                url, 
                headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            )
            # 增加超时限制，防止卡死
            with urllib.request.urlopen(req, timeout=15) as response:
                data = response.read()
                with open(output_path, 'wb') as f:
                    f.write(data)
            print(f"Successfully rendered and saved mermaid chart to: {output_path}")
            return True
        except Exception as e:
            print(f"Failed to render mermaid online (Attempt {attempt+1}/{max_retries}) for path {output_path}: {e}")
            if attempt < max_retries - 1:
                time.sleep(2)
            else:
                return False

def parse_markdown_to_docx(md_path, template_path, output_path, img_base_dir):
    print(f"Loading template from: {template_path}")
    doc = docx.Document(template_path)
    
    # 启用 Word 打开更新域设置，从而允许用户打开时右键提示更新目录
    enable_update_fields_on_open(doc)
    
    # 追加分页符以开始正文
    doc.add_page_break()
    
    # 读取 markdown 文件内容
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    print(f"Read {len(lines)} lines of markdown.")
    
    # 状态机变量
    in_code_block = False
    in_mermaid = False
    code_lines = []
    mermaid_lines = []
    code_lang = ""
    
    in_table = False
    table_rows = []
    
    should_start_writing = False
    
    img_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]*)\)')
    link_pattern = re.compile(r'\[([^\]]*)\]\(([^)]*)\)')
    
    mermaid_counter = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 忽略开头的注释说明及前言，从 "## 摘要" 开始写
        if not should_start_writing:
            if stripped.startswith("## "):
                should_start_writing = True
                print(f"Starting to write report content from line {i}: '{stripped}'")
            else:
                i += 1
                continue
                
        # 1. 解析代码块
        if stripped.startswith("```"):
            if in_mermaid:
                in_mermaid = False
                mermaid_counter += 1
                img_name = f"mermaid_chart_{mermaid_counter}.png"
                img_abs_path = os.path.abspath(os.path.join(img_base_dir, "screenshots", img_name))
                mermaid_content = "".join(mermaid_lines)
                
                print(f"Processing mermaid block {mermaid_counter}...")
                success = download_mermaid_image(mermaid_content, img_abs_path)
                
                title = MERMAID_TITLES.get(mermaid_counter, f"图 4-{mermaid_counter} 系统业务逻辑图")
                
                if success and os.path.exists(img_abs_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_paragraph(p_img, line_spacing=1.0, space_before_pt=6, space_after_pt=4)
                    try:
                        p_img.add_run().add_picture(img_abs_path, width=Cm(13.8))
                    except Exception as e:
                        print(f"Error adding picture {img_abs_path}: {e}")
                        p_img.add_run(f"[图表加载失败: {title}]")
                    
                    p_caption = doc.add_paragraph()
                    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_paragraph(p_caption, line_spacing=1.15, space_before_pt=2, space_after_pt=8)
                    run_caption = p_caption.add_run(title)
                    set_run_font(run_caption, font_name='黑体', font_size_pt=9.5, bold=True)
                else:
                    # 降级打印红字警告
                    p_warn = doc.add_paragraph()
                    p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_warn = p_warn.add_run(f"[警告: 无法在线渲染并生成 Mermaid 图表 - {title}]")
                    set_run_font(run_warn, font_name='宋体', font_size_pt=10, bold=True, color=RGBColor(255, 0, 0))
                    
                mermaid_lines = []
            elif in_code_block:
                in_code_block = False
                write_code_block_to_docx(doc, code_lines, code_lang)
                code_lines = []
            else:
                lang = stripped[3:].strip().lower()
                if lang == "mermaid":
                    in_mermaid = True
                    mermaid_lines = []
                else:
                    in_code_block = True
                    code_lang = lang
                    code_lines = []
            i += 1
            continue
            
        if in_mermaid:
            mermaid_lines.append(line)
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
            
        # 2. 解析表格
        if stripped.startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            
            cols = [col.strip() for col in stripped.split("|")[1:-1]]
            # 过滤掉 markdown 表格对齐指示线行（如 :---, ---, :---: 等）
            is_separator = all(re.match(r'^[\s\-:]*$', c) for c in cols)
            if not is_separator:
                table_rows.append(cols)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                write_table_to_docx(doc, table_rows)
                table_rows = []
                
        # 3. 解析标题
        if stripped.startswith("#"):
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            
            title_text = stripped[level:].strip()
            title_text = link_pattern.sub(r'\1', title_text)
            
            # 设置段落大纲级别，确保即使没有标题样式也能生成目录
            def set_paragraph_outline_level(paragraph, level):
                pPr = paragraph._p.get_or_add_pPr()
                outlineLvl = OxmlElement('w:outlineLvl')
                outlineLvl.set(qn('w:val'), str(level - 1))
                pPr.append(outlineLvl)

            # 自动映射到 Word 标准标题大纲样式，加入中文/英文自适应
            def add_heading_with_fallback(doc, level_num):
                style_en = f'Heading {level_num}'
                style_zh = f'标题 {level_num}'
                p = None
                try:
                    p = doc.add_paragraph(style=style_en)
                except KeyError:
                    try:
                        p = doc.add_paragraph(style=style_zh)
                    except KeyError:
                        p = doc.add_paragraph()
                set_paragraph_outline_level(p, level_num)
                return p

            if level == 2:
                p = add_heading_with_fallback(doc, 1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(title_text)
                set_run_font(run, font_name='黑体', font_size_pt=15, bold=True)
                format_paragraph(p, line_spacing=1.5, space_before_pt=12, space_after_pt=12, first_line_indent_chars=0)
                
                # 摘要、Abstract 居中
                if title_text in ["摘要", "Abstract"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif title_text == "目录":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 特殊处理：目录标题下自动插入目录域
                    p_toc = doc.add_paragraph()
                    p_toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_table_of_contents(p_toc)
            elif level == 3:
                p = add_heading_with_fallback(doc, 2)
                run = p.add_run(title_text)
                set_run_font(run, font_name='黑体', font_size_pt=13, bold=True)
                format_paragraph(p, line_spacing=1.5, space_before_pt=8, space_after_pt=6, first_line_indent_chars=0)
            elif level >= 4:
                p = add_heading_with_fallback(doc, 3)
                run = p.add_run(title_text)
                set_run_font(run, font_name='黑体', font_size_pt=11.5, bold=True)
                format_paragraph(p, line_spacing=1.5, space_before_pt=6, space_after_pt=4, first_line_indent_chars=0)
                
            i += 1
            continue
            
        # 4. 解析图片
        img_match = img_pattern.search(stripped)
        if img_match:
            img_title = img_match.group(1)
            img_rel_path = img_match.group(2)
            # 路径规范化，支持子目录
            img_abs_path = os.path.abspath(os.path.join(img_base_dir, img_rel_path))
            
            print(f"Found image: '{img_title}' at relative: '{img_rel_path}' -> absolute: '{img_abs_path}'")
            
            # 寻找下一行非空行，看看是否是显式的图注（例如 "图 5-1 xxx"）
            caption_text = img_title
            next_idx = i + 1
            has_next_caption = False
            while next_idx < len(lines):
                next_line = lines[next_idx].strip()
                if next_line:
                    if re.match(r'^图\s*\d+-\d+', next_line):
                        caption_text = next_line
                        has_next_caption = True
                    break
                next_idx += 1
            
            if os.path.exists(img_abs_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                format_paragraph(p_img, line_spacing=1.0, space_before_pt=6, space_after_pt=4)
                
                try:
                    p_img.add_run().add_picture(img_abs_path, width=Cm(13.8))
                except Exception as e:
                    print(f"Error adding picture {img_abs_path}: {e}")
                    p_img.add_run(f"[图片加载失败: {img_title}]")
                
                p_caption = doc.add_paragraph()
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                format_paragraph(p_caption, line_spacing=1.15, space_before_pt=2, space_after_pt=8)
                run_caption = p_caption.add_run(caption_text)
                # 使用黑体图注
                set_run_font(run_caption, font_name='黑体', font_size_pt=9.5, bold=True)
            else:
                print(f"WARNING: Image path not found: {img_abs_path}")
                p_warn = doc.add_paragraph()
                p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_warn = p_warn.add_run(f"[警告: 找不到图片 {caption_text}]")
                set_run_font(run_warn, font_name='宋体', font_size_pt=10, bold=True, color=RGBColor(255, 0, 0))
            
            if has_next_caption:
                i = next_idx + 1
            else:
                i += 1
            continue
            
        # 5. 解析列表项
        if stripped.startswith("- ") or stripped.startswith("* "):
            list_text = stripped[2:].strip()
            list_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', list_text)
            list_text = link_pattern.sub(r'\1', list_text)
            
            p = doc.add_paragraph()
            format_paragraph(p, line_spacing=1.5, space_before_pt=0, space_after_pt=3, first_line_indent_chars=0)
            p.paragraph_format.left_indent = Cm(0.75)
            
            bullet_run = p.add_run("• ")
            set_run_font(bullet_run, font_name='Times New Roman', font_size_pt=11.5, bold=True)
            
            run = p.add_run(list_text)
            set_run_font(run, font_name='宋体', font_size_pt=11.5)
            
            i += 1
            continue
            
        if re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^(\d+)\.\s(.*)', stripped)
            num = match.group(1)
            list_text = match.group(2).strip()
            list_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', list_text)
            list_text = link_pattern.sub(r'\1', list_text)
            
            p = doc.add_paragraph()
            format_paragraph(p, line_spacing=1.5, space_before_pt=0, space_after_pt=3, first_line_indent_chars=0)
            p.paragraph_format.left_indent = Cm(0.75)
            
            num_run = p.add_run(f"{num}. ")
            set_run_font(num_run, font_name='Times New Roman', font_size_pt=11.5, bold=True)
            
            run = p.add_run(list_text)
            set_run_font(run, font_name='宋体', font_size_pt=11.5)
            
            i += 1
            continue
            
        # 6. 解析普通段落
        if stripped:
            if stripped == "---":
                doc.add_page_break()
                i += 1
                continue
                
            para_text = stripped
            parts = para_text.split("**")
            
            p = doc.add_paragraph()
            format_paragraph(p, line_spacing=1.5, space_before_pt=0, space_after_pt=4, first_line_indent_chars=2)
            
            for index, part in enumerate(parts):
                if not part:
                    continue
                is_bold = (index % 2 == 1)
                cleaned_part = link_pattern.sub(r'\1', part)
                run = p.add_run(cleaned_part)
                set_run_font(run, font_name='宋体', font_size_pt=11.5, bold=is_bold)
                
        i += 1
        
    # --- 循环结束后的收尾逻辑 ---
    if in_table and table_rows:
        print("Flushing final leftover table from buffer.")
        write_table_to_docx(doc, table_rows)
        
    if in_code_block and code_lines:
        print("Flushing final leftover code block from buffer.")
        write_code_block_to_docx(doc, code_lines, code_lang)
        
    # 保存文档
    print(f"Saving final report to: {output_path}")
    doc.save(output_path)
    print("Success!")

if __name__ == "__main__":
    md = "d:\\project\\equipment-management\\docs\\4-tasks\\features\\TASK-019-write-web-course-report\\course_design_report.md"
    template = "d:\\project\\equipment-management\\tmp\\docs\\23WEB课程设计报告-converted.docx"
    output = "d:\\project\\equipment-management\\docs\\4-tasks\\features\\TASK-019-write-web-course-report\\course_design_report_final.docx"
    img_dir = "d:\\project\\equipment-management\\docs\\4-tasks\\features\\TASK-019-write-web-course-report"
    
    parse_markdown_to_docx(md, template, output, img_dir)
